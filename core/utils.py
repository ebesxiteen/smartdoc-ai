"""
Centralized utility functions for RAG application.
Includes: PDF processing, chat handling, vectorstore management, and text cleaning.
"""

from datetime import datetime, timezone
import hashlib
import uuid
from typing import Any, Dict, List, Optional, Set, Tuple, Callable, cast
import logging

import streamlit as st
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings

import core.configs as cfg
from langchain_core.messages import BaseMessage
from pathlib import Path
import re
from langdetect import (  # pyright: ignore[reportMissingTypeStubs]
    detect,  # pyright: ignore[reportUnknownVariableType]
)
import psutil
import platform

logger = logging.getLogger(__name__)


# ============================================================================
# HARDWARE DETECTION UTILITY
# ============================================================================


@st.cache_data(show_spinner=False)
def get_system_hardware_info() -> Dict[str, Any]:
    """
    Detect and gather system hardware capabilities including CPU, RAM, and GPU.

    This function probes the system for hardware information to enable dynamic
    UI warnings and RAG parameter optimization based on available resources.
    GPU detection attempts to use PyTorch/CUDA if available, with graceful fallback
    to CPU-only mode if GPU detection fails.

    Returns:
        Dict[str, Any]: Hardware capabilities dictionary containing:
            - os (str): Operating system name (e.g., "Linux", "Darwin", "Windows")
            - cpu_cores (int): Number of logical CPU cores
            - ram_gb (float): Total RAM in gigabytes
            - gpu_name (str): GPU model name or "None/Integrated" if not available
            - vram_gb (float): Total GPU VRAM in gigabytes (0.0 if no GPU)

    Example:
        >>> hw_info = get_system_hardware_info()
        >>> if hw_info["vram_gb"] > 0:
        ...     print(f"GPU detected: {hw_info['gpu_name']}")
    """
    # OS & CPU & RAM
    os_name = platform.system()
    cpu_cores = psutil.cpu_count(logical=True)
    total_ram_gb = round(psutil.virtual_memory().total / (1024**3), 1)

    # GPU
    gpu_name = "None/Integrated"
    total_vram_gb = 0.0

    try:
        import torch

        if torch.cuda.is_available():
            gpu_name = torch.cuda.get_device_name(0)
            # torch limits might not reflect exact system VRAM,
            # but usually get_device_properties gives total_memory in bytes
            props = getattr(torch.cuda, "get_device_properties")(0)
            total_memory_bytes = getattr(props, "total_memory", 0)
            total_vram_gb = round(int(total_memory_bytes) / (1024**3), 1)
    except Exception as e:
        debug_log("WARNING", message=f"GPU detection failed: {e}")

    return {
        "os": os_name,
        "cpu_cores": cpu_cores,
        "ram_gb": total_ram_gb,
        "gpu_name": gpu_name,
        "vram_gb": total_vram_gb,
    }


def get_default_notebook_settings() -> Dict[str, Any]:
    """
    Get default notebook configuration settings for RAG pipeline initialization.

    Returns a dictionary of default settings covering retrieval parameters (K value,
    score thresholds, chunk sizes), LLM configuration (model name, temperature,
    context window), and hybrid search weights (semantic vs. keyword balance).

    Returns:
        Dict[str, Any]: Default settings dictionary with keys:
            - rag_final_context_k (int): Number of top chunks to retrieve
            - rag_rerank_top_n (int): Number of chunks to rerank with Cross-Encoder
            - rag_retrieval_min_results (int): Minimum guaranteed results
            - rag_retrieval_score_threshold (float): Similarity score threshold
            - rag_max_chunk_len (int): Maximum chunk size in characters
            - rag_chunk_overlap (int): Overlap between consecutive chunks
            - rag_max_ctx_len (int): Maximum context length for LLM
            - max_msg_history (int): Maximum chat history messages to retain
            - llm_model_name (str): Ollama model name
            - llm_num_ctx (int): LLM context window size
            - llm_avg_temp (float): LLM temperature (0.0-2.0)
            - personal_ctx (None): Custom instructions placeholder
            - weight_semantic (float): Weight for semantic (0.0-1.0)
            - weight_bm25 (float): Weight for BM25 keyword search (0.0-1.0)
            - self_rag_max_depth (int): Maximum recursion depth for Self-RAG repairs
            - self_rag_candidates (int): Max candidate sub-queries per hop
            - self_rag_max_retries_per_hop (int): Max retries per hop to prevent oscillation
            - self_rag_threshold_issup (float): Threshold for "Is Supportive" in Self-RAG
            - self_rag_threshold_isrel (float): Threshold for "Is Relevant" in Self-RAG
            - self_rag_threshold_isuse (float): Threshold for "Is Useful" in Self-RAG

    Note:
        All values are loaded from core/configs.py. Modify settings there to change defaults.
    """
    return {
        "rag_final_context_k": cfg.RAG_FINAL_CONTEXT_K,
        "rag_rerank_top_n": cfg.RAG_RERANK_TOP_N,
        "rag_retrieval_min_results": cfg.RAG_RETRIEVAL_MIN_RESULTS,
        "rag_retrieval_score_threshold": cfg.RAG_RETRIEVAL_SCORE_THRESHOLD,
        "rag_max_chunk_len": cfg.RAG_MAX_CHUNK_LEN,
        "rag_chunk_overlap": cfg.RAG_CHUNK_OVERLAP,
        "rag_max_ctx_len": cfg.RAG_MAX_CTX_LEN,
        "max_msg_history": cfg.MAX_MSG_HISTORY,
        "llm_model_name": cfg.LLM_MODEL_NAME,
        "llm_num_ctx": cfg.LLM_NUM_CTX,
        "llm_avg_temp": cfg.LLM_AVG_TEMP,
        "personal_ctx": None,
        "weight_semantic": cfg.WEIGHT_SEMANTIC,
        "weight_bm25": cfg.WEIGHT_BM25,
        "self_rag_max_depth": cfg.SELF_RAG_MAX_DEPTH,
        "self_rag_candidates": cfg.SELF_RAG_CANDIDATES,
        "self_rag_max_retries_per_hop": cfg.SELF_RAG_MAX_RETRIES_PER_HOP,
        "self_rag_threshold_issup": cfg.SELF_RAG_THRESHOLD_ISSUP,
        "self_rag_threshold_isrel": cfg.SELF_RAG_THRESHOLD_ISREL,
        "self_rag_threshold_isuse": cfg.SELF_RAG_THRESHOLD_ISUSE,
        "co_rag_max_retries": cfg.CO_RAG_MAX_RETRIES,
    }


@st.cache_data(show_spinner=False)
def load_notebook_settings(
    notebook_id: Optional[str],
) -> Dict[str, Any]:
    """
    Load notebook-specific settings from the database, with fallback to defaults.

    Retrieves persisted settings for a specific notebook from the database.
    If the notebook ID is None or notebook has no custom settings, returns
    default settings from get_default_notebook_settings().

    Args:
        notebook_id: The UUID of the notebook to load settings for.
                    If None, default settings are returned.

    Returns:
        Dict[str, Any]: Complete settings dictionary for the notebook.
                       If notebook has custom settings, those are merged with defaults.

    Note:
        This is an internal function used to ensure all RAG chains always have
        complete settings dictionaries, even if the notebook doesn't have
        explicitly saved settings.
    """
    defaults = get_default_notebook_settings()
    if not notebook_id:
        return defaults

    from middlewares.db_middleware import get_notebook_settings

    settings = get_notebook_settings(notebook_id)
    if not settings:
        return defaults

    # Use `or defaults[key]` so that DB NULL values fall back to config constants.
    # dict.get(key, default) only uses `default` when the key is absent; it returns
    # None when the key exists with a NULL value — which would silently break callers
    # that do arithmetic with the result (e.g., top_k=None crashes rerank_with_cross_encoder).
    return {
        "rag_final_context_k": settings.get("rag_final_context_k")
        or defaults["rag_final_context_k"],
        "rag_rerank_top_n": settings.get("rag_rerank_top_n")
        or defaults["rag_rerank_top_n"],
        "rag_retrieval_min_results": settings.get("rag_retrieval_min_results")
        if settings.get("rag_retrieval_min_results") is not None
        else defaults["rag_retrieval_min_results"],
        "rag_retrieval_score_threshold": settings.get("rag_retrieval_score_threshold")
        or defaults["rag_retrieval_score_threshold"],
        "rag_max_chunk_len": settings.get("rag_max_chunk_len")
        or defaults["rag_max_chunk_len"],
        "rag_chunk_overlap": settings.get("rag_chunk_overlap")
        if settings.get("rag_chunk_overlap") is not None
        else defaults["rag_chunk_overlap"],
        "rag_max_ctx_len": settings.get("rag_max_ctx_len")
        or defaults["rag_max_ctx_len"],
        "max_msg_history": settings.get("max_msg_history")
        or defaults["max_msg_history"],
        "llm_model_name": settings.get("llm_model_name") or defaults["llm_model_name"],
        "llm_num_ctx": settings.get("llm_num_ctx") or defaults["llm_num_ctx"],
        # DB column is `llm_avg_temp`; read with that key and fall back to config default.
        "llm_avg_temp": settings.get("llm_avg_temp") or defaults["llm_avg_temp"],
        "personal_ctx": settings.get("personal_ctx"),
        "weight_semantic": settings.get("weight_semantic")
        if settings.get("weight_semantic") is not None
        else defaults["weight_semantic"],
        "weight_bm25": settings.get("weight_bm25")
        if settings.get("weight_bm25") is not None
        else defaults["weight_bm25"],
        "self_rag_max_depth": settings.get("self_rag_max_depth")
        if settings.get("self_rag_max_depth") is not None
        else defaults["self_rag_max_depth"],
        "self_rag_candidates": settings.get("self_rag_candidates")
        or defaults["self_rag_candidates"],
        "self_rag_max_retries_per_hop": settings.get("self_rag_max_retries_per_hop")
        if settings.get("self_rag_max_retries_per_hop") is not None
        else defaults["self_rag_max_retries_per_hop"],
        "self_rag_threshold_issup": settings.get("self_rag_threshold_issup")
        if settings.get("self_rag_threshold_issup") is not None
        else defaults["self_rag_threshold_issup"],
        "self_rag_threshold_isrel": settings.get("self_rag_threshold_isrel")
        if settings.get("self_rag_threshold_isrel") is not None
        else defaults["self_rag_threshold_isrel"],
        "self_rag_threshold_isuse": settings.get("self_rag_threshold_isuse")
        if settings.get("self_rag_threshold_isuse") is not None
        else defaults["self_rag_threshold_isuse"],
        "co_rag_max_retries": settings.get("co_rag_max_retries")
        if settings.get("co_rag_max_retries") is not None
        else defaults["co_rag_max_retries"],
    }


def generate_fallback_answer(
    user_query: str,
    llm_chain: Any,
    chat_history: Optional[List[Any]] = None,
    print_debug: bool = False,
) -> str:
    """
    Generate LLM-based fallback answer when no relevant documents found.

    Creates a context-aware fallback response using the LLM instead of hard-coded messages,
    providing better user experience with dynamic, thoughtful answers.

    Args:
        user_query: The original user question
        llm_chain: LLM chain for answer generation
        chat_history: Optional formatted conversation history for context
        print_debug: Whether to print debug logs

    Returns:
        str: LLM-generated fallback answer (context-aware, dynamic)
    """
    try:
        from langchain_core.prompts import ChatPromptTemplate

        # Estimate query complexity based on length and keyword presence
        query_length = len(user_query.split())
        complex_keywords = {
            "how to",
            "why is",
            "explain",
            "compare",
            "analyze",
            "what is the",
        }
        is_complex = query_length > 10 or any(
            keyword in user_query.lower() for keyword in complex_keywords
        )
        if is_complex:
            query_complexity = "complex"
        elif query_length < 5:
            query_complexity = "simple"
        else:
            query_complexity = "medium"

        # Format chat history if available (supports both dict and LangChain message objects)
        formatted_history = ""
        if chat_history:
            history_lines: List[str] = []
            for msg in chat_history[-3:]:  # Last 3 messages
                if hasattr(msg, "get"):
                    role = (
                        "User"
                        if msg.get("role", msg.get("type", "")) in ("user", "human")
                        else "Assistant"
                    )
                    content = msg.get("content", "")
                else:
                    # LangChain BaseMessage (HumanMessage / AIMessage)
                    msg_type = getattr(msg, "type", "")
                    role = "User" if msg_type == "human" else "Assistant"
                    content = getattr(msg, "content", "")
                history_lines.append(f"{role}: {content}")
            formatted_history = "\n".join(history_lines)

        # Generate fallback prompt
        fallback_prompt = cfg.generate_general_knowledge_fallback_prompt(
            user_query=user_query,
            chat_history=formatted_history if formatted_history else None,
            query_complexity=query_complexity,
        )

        # Create prompt template and chain
        prompt = ChatPromptTemplate.from_template(fallback_prompt)
        chain = prompt | llm_chain

        # Invoke LLM for fallback answer
        response = chain.invoke({})
        fallback_answer = str(response).strip()

        if print_debug:
            debug_log(
                "INFO",
                "💭",
                f"Generated dynamic fallback answer ({len(fallback_answer)} chars, complexity: {query_complexity})",
            )

        return fallback_answer

    except Exception as e:
        if print_debug:
            debug_log("WARNING", "⚠️", f"Fallback generation failed: {str(e)[:100]}")
        # Fallback to hard-coded message if LLM fails
        return cfg.NOT_FOUND_ANSWER_FALL_BACK


# ============================================================================
# TEXT UTILITIES
# ============================================================================


def clean_spaces(text: str) -> str:
    """
    Normalize whitespace in text by removing leading/trailing spaces and
    collapsing multiple internal spaces into single spaces.

    This is useful for cleaning text extracted from PDFs or other sources
    that may have inconsistent whitespace formatting.

    Args:
        text: Input string to normalize.

    Returns:
        str: Text with normalized whitespace. Returns empty string if input is falsy.

    Example:
        >>> clean_spaces("  Hello    world  ")
        'Hello world'
    """
    if not text:
        return ""
    # .split() without arguments splits by any whitespace (space, \n, \t)
    # ' '.join(...) puts exactly one space between each word
    return " ".join(text.split())


# ============================================================================
# DEBUG LOGGING UTILITIES & CONSTANTS
# ============================================================================


def print_breaker() -> None:
    """
    Print a visual separator line in logs to demarcate logical sections.

    Outputs a 70-character wide line (━) to visually divide log sections,
    improving readability when scanning terminal output.
    """
    separator = "━" * 70
    logger.info(separator)


def debug_log(
    log_type: str = "INFO",
    emoji: Optional[str] = None,
    message: str = "",
    color: Optional[str] = None,
    show_metadata: bool = True,
) -> None:
    """
    Enhanced structured logging with automatic categorization and ANSI color support.

    This logger provides consistent, visually distinct log messages with automatic
    color mapping, caller metadata extraction, and support for structured log categories.

    Args:
        log_type: Log category or severity type. Can be:
                 - Standard type: "INFO", "WARNING", "ERROR", "SUCCESS"
                 - Structured category: Any key from LOG_CATEGORIES dict
                 Default: "INFO"
        emoji: Emoji to prefix the log message. If None and log_type is a key
               in LOG_CATEGORIES, the emoji from that category is used.
               Default: None
        message: The main log message content.
                Default: ""
        color: Optional ANSI color code override (e.g., "31" for red).
               If not provided, color is determined by log_type.
               Default: None
        show_metadata: Whether to include caller filename and line number in output.
                      Default: True

    Example:
        >>> debug_log("SUCCESS", "✅", "File loaded successfully")
        >>> debug_log("KEYWORD", message="Using BM25 keyword search")
        >>> debug_log("WARNING", message="Threshold not met, using fallback")
    """
    import inspect
    import os

    # 1. Define Default Color Map
    # ANSI color codes: 31=Red, 33=Yellow, 36=Cyan, 32=Green
    COLOR_MAP = {
        "ERROR": "31",
        "WARNING": "33",
        "INFO": "36",
        "SUCCESS": "32",
    }

    # 2. Normalize type and determine emoji + color
    normalized_type = (log_type or "INFO").upper()

    # Check if it's a structured category
    if normalized_type in cfg.LOG_CATEGORIES:
        category = cfg.LOG_CATEGORIES[normalized_type]
        used_emoji = emoji or category["emoji"]
        used_type = category["type"]
        final_color = color if color else COLOR_MAP.get(used_type, "36")
    else:
        used_emoji = emoji
        used_type = normalized_type
        final_color = color if color else COLOR_MAP.get(normalized_type, "36")

    # 3. Capture caller metadata (if enabled)
    metadata_str = ""
    if show_metadata:
        frame = inspect.stack()[1]
        filename = os.path.basename(frame.filename)
        lineno = frame.lineno
        metadata_str = f"\033[{final_color}m[{filename}:{lineno}]\033[0m"

    # 4. Format the output
    # \u2009 is a thin space to prevent emoji overlap
    prefix = f"{used_emoji}\u2009" if used_emoji else ""
    if metadata_str:
        printed_message = f"{prefix}{metadata_str} {message}"
    else:
        printed_message = f"{prefix}{message}"

    # 5. Route to logger based on type
    if used_type == "ERROR":
        logger.error(printed_message)
    elif used_type == "WARNING":
        logger.warning(printed_message)
    elif used_type == "SUCCESS":
        logger.info(printed_message)
    else:
        logger.info(printed_message)


# ============================================================================
# PDF PROCESSING UTILITIES
# ============================================================================


def load_and_chunk_file(
    file_path: str,
    file_type: str,
    chunk_size: int = cfg.RAG_MAX_CHUNK_LEN,
    chunk_overlap: int = cfg.RAG_CHUNK_OVERLAP,
    print_debug: bool = False,
    progress_callback: Optional[Callable[[str], None]] = None,
) -> List[Document]:
    """
    Load a file (PDF or DOCX) and split it into overlapping text chunks.

    This function extracts text from PDF or DOCX files and uses a recursive
    text splitter to create overlapping chunks suitable for embedding and RAG.
    Progress updates can be provided via callback for UI integration.

    Args:
        file_path: Absolute path to the file to load.
        file_type: File format - either "pdf" or "docx".
        chunk_size: Maximum characters per chunk. Default from configs.py.
        chunk_overlap: Character overlap between consecutive chunks. Default from configs.py.
        print_debug: If True, log detailed processing information. Default False.
        progress_callback: Optional callable to receive progress messages (e.g., for UI progress bar).

    Returns:
        List[Document]: LangChain Document objects with chunked text content and metadata.
                       Each document contains:
                       - page_content: Chunk text
                       - metadata: {"source": file_path, "page": page_number, "document": filename}

    Raises:
        ValueError: If file_type is neither "pdf" nor "docx".
        Exception: Various exceptions from PyMuPDF or python-docx if file parsing fails.

    Example:
        >>> chunks = load_and_chunk_file(
        ...     "document.pdf",
        ...     "pdf",
        ...     chunk_size=1000,
        ...     print_debug=True
        ... )
        >>> print(f"Created {len(chunks)} chunks")
    """
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    # Step 1: Load the Document
    documents: List[Document] = []
    if file_type == "pdf":
        from langchain_community.document_loaders import PyMuPDFLoader

        loader = PyMuPDFLoader(file_path)
        for i, page in enumerate(loader.lazy_load()):
            documents.append(page)
            if progress_callback:
                progress_callback(f"Reading page {i + 1}...")
    elif file_type == "docx":
        import docx

        if progress_callback:
            progress_callback("Reading DOCX file...")
        doc = docx.Document(file_path)
        full_text: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        text = "\n".join(full_text)
        if text.strip():
            documents = [Document(page_content=text, metadata={"source": file_path})]
    else:
        raise ValueError(f"Unsupported file type for loading: {file_type}")

    if print_debug:
        debug_log(
            "SUCCESS",
            message=f"Loaded {file_type.upper()}: {file_path}\n      {len(documents)} pages | {sum(len(d.page_content) for d in documents)} chars",
        )

    # Step 2: Split into chunks
    if progress_callback:
        progress_callback("Chunking text...")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=[
            "\n\n",
            "\n",
            " ",
            "",
        ],  # Split by paragraph, then line, then space, then character
    )

    chunks = text_splitter.split_documents(documents)

    if print_debug:
        avg_chunk_size = (
            sum(len(chunk.page_content) for chunk in chunks) / len(chunks)
            if chunks
            else 0
        )
        debug_log(
            "COMPLETE",
            message=f"Chunking complete: {len(chunks)} chunks\n      Avg chunk size: {avg_chunk_size:.0f} chars",
        )

    return chunks


def hash_file_content(file_bytes: bytes) -> str:
    """
    Calculate MD5 hash of file content for duplicate detection.

    Args:
        file_bytes: Raw file content as bytes.

    Returns:
        str: Hex-encoded MD5 hash of the file content.

    Note:
        MD5 is used for duplicate detection only, not security. SHA256 could
        be used for enhanced security if needed in future.
    """
    return hashlib.md5(file_bytes).hexdigest()


def detect_file_type(file_bytes: bytes) -> str:
    """
    Detect file type based on magic number (file signature bytes).

    Args:
        file_bytes: First few bytes of the file content.

    Returns:
        str: Detected file type - either "pdf" or "docx".

    Raises:
        ValueError: If file type is neither PDF nor DOCX.
    """
    if file_bytes.startswith(b"%PDF"):
        return "pdf"
    elif (
        file_bytes.startswith(b"PK\x03\x04")
        or file_bytes.startswith(b"PK\x05\x06")
        or file_bytes.startswith(b"PK\x07\x08")
    ):
        return "docx"
    else:
        raise ValueError(
            "Unsupported file format. Only PDF and Word (docx) are supported."
        )


def check_file_already_exists(file_hash: str) -> Optional[Dict[str, Any]]:
    """
    Check if a file with the given hash was already uploaded to any notebook.

    Args:
        file_hash: MD5 hash of file content from hash_file_content().

    Returns:
        Dict[str, Any]: Source document info if found, None if not found.
    """
    from db.crud import get_source_by_hash

    return get_source_by_hash(file_hash)


def check_file_already_exists_in_notebook(
    file_hash: str, notebook_id: str
) -> Optional[Dict[str, Any]]:
    """
    Check if a file with the given hash was already uploaded to a specific notebook.

    Args:
        file_hash: MD5 hash of file content from hash_file_content().
        notebook_id: UUID of the notebook to check within.

    Returns:
        Dict[str, Any]: Source document info if found in this notebook, None if not found.
    """
    from db.crud import get_source_by_hash_and_notebook

    return get_source_by_hash_and_notebook(file_hash, notebook_id)


def chunk_and_process_file(
    file_path: str,
    file_type: str,
    filename: str,
    print_debug: bool = False,
    progress_callback: Optional[Callable[[str], None]] = None,
    chunk_size: int = cfg.RAG_MAX_CHUNK_LEN,
    chunk_overlap: int = cfg.RAG_CHUNK_OVERLAP,
) -> Tuple[List[Document], int]:
    """
    Load a file, split into chunks, and enrich with metadata.

    Wraps load_and_chunk_file() to add the original filename to chunk metadata
    for source attribution in RAG responses.

    Args:
        file_path: Absolute path to the file.
        file_type: "pdf" or "docx".
        filename: Display name of the file (for metadata).
        print_debug: Enable debug logging.
        progress_callback: Optional progress update callback.
        chunk_size: Maximum chunk size in characters.
        chunk_overlap: Character overlap between chunks.

    Returns:
        Tuple[List[Document], int]: Tuple of:
            - List of Document objects with metadata
            - Count of created chunks
    """
    if print_debug:
        debug_log("INFO", "📄", f"Loading and chunking {file_type.upper()}: {filename}")
    chunks = load_and_chunk_file(
        file_path,
        file_type,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        print_debug=print_debug,
        progress_callback=progress_callback,
    )

    # Add document name to metadata
    for chunk in chunks:
        chunk.metadata["document"] = filename  # type: ignore[index]

    return chunks, len(chunks)


def create_vectorstore_from_chunks(
    chunks: List[Document],
    embeddings: HuggingFaceEmbeddings,
    print_debug: bool = False,
    progress_callback: Optional[Callable[[str], None]] = None,
) -> FAISS:
    """
    Create a FAISS vector database from document chunks.

    Args:
        chunks: List of Document objects with text content to embed and index.
        embeddings: HuggingFaceEmbeddings instance for vectorization.
        print_debug: Enable debug logging.
        progress_callback: Optional callback for status.

    Returns:
        FAISS: Vector store instance ready for similarity search.
    """
    if print_debug:
        debug_log("EMBED", message=f"Creating vectorstore from {len(chunks)} chunks...")

    if not progress_callback:
        return FAISS.from_documents(chunks, embeddings)

    # Batch process for progress updates
    batch_size = max(10, len(chunks) // 10)  # Update ~10 times ideally
    vectorstore = None

    for i in range(0, len(chunks), batch_size):
        batch = chunks[i : i + batch_size]
        progress_msg = f"Embedding chunks {i + 1} to {min(i + batch_size, len(chunks))} of {len(chunks)}..."
        progress_callback(progress_msg)

        if vectorstore is None:
            vectorstore = FAISS.from_documents(batch, embeddings)
        else:
            vectorstore.add_documents(batch)

    if vectorstore is None:
        raise ValueError("Cannot create vectorstore from empty chunks list.")

    return vectorstore


def merge_vectorstores(
    existing_vectorstore: Optional[FAISS],
    new_vectorstore: FAISS,
    print_debug: bool = False,
) -> FAISS:
    """
    Merge a new vectorstore into an existing one, or return the new one if none exists.

    Used when multiple sources are added to a notebook to maintain a consolidated
    vectorstore for efficient similarity search across all sources.

    Args:
        existing_vectorstore: Current FAISS vectorstore or None.
        new_vectorstore: New FAISS vectorstore to merge in.
        print_debug: Enable debug logging.

    Returns:
        FAISS: Merged vectorstore combining both inputs, or new_vectorstore if no existing one.
    """
    if existing_vectorstore is None:
        if print_debug:
            debug_log("TASK_START", message="No existing vectorstore - using new one")
        return new_vectorstore

    if print_debug:
        debug_log("MERGED", message="Vectorstores merged successfully")
    existing_vectorstore.merge_from(new_vectorstore)
    return existing_vectorstore


def save_source_to_database(
    notebook_id: str,
    filename: str,
    file_type: str,
    file_hash: str,
    summary: str,
    suggested_questions: List[str],
    vectorstore_path: str,
    source_id: Optional[str] = None,
    print_debug: bool = False,
) -> str:
    """
    Save source file metadata and vectorstore location to the database.

    Records a new source document in the notebook, linking it to its FAISS vectorstore.
    Enables tracking of document sources for citation and content refresh.

    Args:
        notebook_id: UUID of the notebook this source belongs to.
        filename: Original filename displayed to users.
        file_type: "pdf" or "docx".
        file_hash: MD5 hash of file content for duplicate detection.
        summary: AI-generated summary of the document.
        suggested_questions: List of auto-generated discussion questions.
        vectorstore_path: Path to the saved FAISS vectorstore directory.
        source_id: Optional pre-generated UUID. If None, one is created.
        print_debug: Enable debug logging.

    Returns:
        str: The source_id (UUID) assigned to this source.

    Note:
        The vectorstore_path is typically computed via get_source_vectorstore_dir().
    """
    from middlewares import db_middleware as db

    if print_debug:
        debug_log("SAVED", message=f"Saving source to database: {filename}")

    if source_id is None:
        source_id = str(uuid.uuid4())

    saved_source_id = db.add_source(
        notebook_id=notebook_id,
        file_name=filename,
        file_type=file_type,
        file_path=vectorstore_path,
        file_hash=file_hash,
        summary=summary,
        suggested_questions=suggested_questions,
        source_id=source_id,
    )

    return saved_source_id


# ============================================================================
# CHAT PROCESSING UTILITIES
# ============================================================================


def process_user_query(
    query: str,
    rag_chain: Any,
    vectorstore: FAISS,
    chat_history: Optional[List[Dict[str, Any]]] = None,
    print_debug: bool = False,
    notebook_id: Optional[str] = None,
) -> tuple[str, List[Dict[str, Any]], bool, List[str], Optional[float]]:
    """
    Process a user query through the Self-RAG pipeline to generate an answer with source citations.

    Handles three query scenarios:
    1. Greetings/General knowledge: Answer without document retrieval
    2. Follow-up questions: Use chat history to rephrase query for better retrieval
    3. Standalone questions: Multi-hop document retrieval with quality-based repair

    Args:
        query: The user's question.
        rag_chain: The instantiated RAG chain from create_history_aware_rag_chain() (uses Self-RAG internally).
        vectorstore: FAISS vectorstore for document retrieval.
        chat_history: Optional list of previous Q&A exchanges for context.
        print_debug: Enable detailed debug logging.
        notebook_id: UUID of the notebook for loading settings.

    Returns:
        Tuple[str, List[Dict[str, Any]], bool, List[str]]:
            - answer (str): Clean LLM response (with status tags removed)
            - sources (List[Dict]): Citation sources with keys: document, page, content
            - found_answer (bool): Whether relevant context was found (affects UI display)
            - reasoning_trace (List[str]): Self-RAG execution trace for transparency (optional UI display)

    Note:
        Self-RAG now handles all answer generation internally through the orchestrated pipeline.
        Answers are tagged with [STATUS: DOC_ANSWER], [STATUS: DOC_MISSING], or [STATUS: GENERAL].
        These tags are stripped before returning to maintain clean output.
        Reasoning trace is extracted from st.session_state["self_rag_metadata"] for UI transparency.
    """
    import streamlit as st

    if print_debug:
        query_preview = f"{query[:60]}..." if len(query) > 60 else query
        debug_log("QUERY", message=f"Processing: {query_preview}")

    # Step 1: Check if this is a greeting/general question
    is_greeting_query = is_greeting(query)
    if print_debug and is_greeting_query:
        debug_log("DEBUG", message="Detected greeting → using general knowledge")

    settings = load_notebook_settings(notebook_id)

    # Step 2: Prepare inputs for the RAG chain
    # Build chain input dict with proper structure for history-aware chain
    # Using a list for retrieved_docs so we can safely extract it back from the chain
    retrieved_docs: List[Document] = []
    chain_input_dict: Dict[str, Any] = {
        "input": query,
        "question": query,
        "__retrieved_docs__": retrieved_docs,
        "is_greeting": is_greeting_query,
    }

    if chat_history:
        formatted_history = format_chat_history_for_rephrase(
            chat_history, max_messages=int(settings["max_msg_history"])
        )
        chain_input_dict["chat_history"] = formatted_history
        if print_debug:
            print_breaker()
            debug_log(
                "HISTORY", message=f"Chat history: {len(formatted_history)} messages"
            )
            for msg in formatted_history:
                role_icon = "👤" if msg.type == "human" else "🤖"
                content_str = str(msg.content)  # type: ignore
                content_preview = content_str.replace("\n", " ")
                content_preview = (
                    content_preview[:70] + "..."
                    if len(content_preview) > 70
                    else content_preview
                )
                debug_log("ITEM", emoji=role_icon, message=content_preview)
            print_breaker()

    # Step 3: Generate answer through Self-RAG chain
    if print_debug:
        if is_greeting_query:
            debug_log("CHAIN", message="Invoking Self-RAG [General Knowledge Mode]...")
        else:
            debug_log(
                "CHAIN", message="Invoking Self-RAG [Multi-Hop Retrieval Mode]..."
            )

    # Initialize empty reasoning trace (will be populated by Self-RAG)
    reasoning_trace: List[str] = []

    try:
        # Invoke chain with proper input structure
        # The Self-RAG chain expects dict with "input" and optional "chat_history"
        # It internally calls core.self_rag.self_rag_query() orchestrator
        answer = rag_chain.invoke(chain_input_dict)

        # Capture reasoning trace from session state (populated by Self-RAG wrapper)
        if "self_rag_metadata" in st.session_state:
            reasoning_trace = st.session_state.self_rag_metadata.get(
                "reasoning_trace", []
            )
            if print_debug:
                debug_log(
                    "INFO",
                    "📋",
                    f"Captured Self-RAG reasoning trace ({len(reasoning_trace)} steps)",
                )

    except Exception as e:
        debug_log("ERROR", message=f"RAG chain failed: {e}")
        answer = f"[STATUS: DOC_MISSING]\nI encountered an error processing your query: {str(e)}"

    if print_debug:
        debug_log("RESPONSE", message=f"LLM response generated ({len(answer)} chars)")

        # Log the actual answer cleanly
        print_breaker()
        debug_log("LLM_INIT", message="LLM Answer:")

        # Split by newlines so it aligns cleanly in the terminal
        for idx, line in enumerate(answer.split("\n")):
            if line.strip():
                # To prevent overwhelming logs, limit to first 10 lines max or 1000 chars
                if idx > 9 or len(line) > 150:
                    preview = line[:150] + "..." if len(line) > 150 else line
                    debug_log("INFO", message=f"{preview}")
                else:
                    debug_log("INFO", message=f"{line}")
        print_breaker()

    # Step 4: Parse [STATUS: DOC_ANSWER/DOC_MISSING/GENERAL] tag from answer
    found_answer = True  # Default to True
    answer_clean = answer
    is_general_answer = is_greeting_query

    # Check for tags using robust regex.
    # IMPORTANT: DOC_ANSWER must explicitly reset is_general_answer to False.
    # Without this, a raw query that matches a greeting pattern (e.g., "Thanks! Tell me
    # about revenue") would leave is_general_answer=True even when Self-RAG returns a
    # document-grounded answer, silently suppressing the source citations in the UI.
    if re.search(r"\[STATUS:\s*DOC_ANSWER\]", answer, flags=re.IGNORECASE):
        found_answer = True
        is_general_answer = (
            False  # DOC_ANSWER always means document-grounded; show sources
        )
    elif re.search(r"\[STATUS:\s*DOC_MISSING\]", answer, flags=re.IGNORECASE):
        found_answer = False
        is_general_answer = False
    elif re.search(r"\[STATUS:\s*GENERAL\]", answer, flags=re.IGNORECASE):
        found_answer = False
        is_general_answer = True

    # Strip out any tags and clean up
    answer_clean = re.sub(
        r"\[STATUS:\s*(?:DOC_ANSWER|DOC_MISSING|GENERAL)\]",
        "",
        answer,
        flags=re.IGNORECASE,
    ).strip()

    # Prevent 'Content cannot be empty' DB error if LLM only returned a tag in rare cases
    if not answer_clean:
        if print_debug:
            debug_log(
                "INFO",
                "⚠️",
                "LLM returned only tags - generating dynamic fallback answer",
            )

        # Create LLM chain for fallback generation if needed
        try:
            from langchain_ollama import OllamaLLM

            settings = load_notebook_settings(notebook_id)
            llm = OllamaLLM(
                model=settings.get("llm_model_name", cfg.LLM_MODEL_NAME),
                base_url=cfg.LLM_BASE_URL,
                temperature=settings.get("llm_avg_temp", cfg.LLM_AVG_TEMP),
            )
            answer_clean = generate_fallback_answer(
                query,
                llm,
                chat_history,
                print_debug,
            )
        except Exception as e:
            if print_debug:
                debug_log("WARNING", "⚠️", f"Fallback generation failed: {str(e)[:100]}")
            answer_clean = cfg.NOT_FOUND_ANSWER_FALL_BACK

    if print_debug:
        if found_answer and not is_general_answer:
            debug_log(
                "INFO", "📍", "LLM found relevant context - sources will be displayed"
            )
        elif is_general_answer:
            debug_log(
                "INFO", "💬", "LLM answered from general knowledge - no sources needed"
            )
        else:
            debug_log(
                "INFO",
                "⚠️",
                "LLM did not find relevant context - sources will be hidden",
            )

    # Step 5: Retrieve source documents for display (only if not a general question)
    sources: List[Dict[str, Any]] = []

    if not is_general_answer:
        if print_debug:
            debug_log("INFO", "📎", "Gathering exact retrieved sources for display...")

        # Prefer the sources captured by the Self-RAG pipeline and stored in session state.
        # chain_input_dict["__retrieved_docs__"] is never populated by self_rag_wrapper()
        # because the wrapper only writes to st.session_state.self_rag_metadata["sources"].
        # Using that session state value guarantees we show the exact chunks the winning
        # candidate was grounded in, not a fresh similarity search on the raw query.
        self_rag_sources: List[Dict[str, Any]] = []
        if (
            getattr(st, "session_state", None)
            and "self_rag_metadata" in st.session_state
        ):
            self_rag_sources = st.session_state.self_rag_metadata.get("sources", [])

        if self_rag_sources:
            sources = self_rag_sources
            if print_debug:
                debug_log(
                    "INFO",
                    "📎",
                    f"Using {len(sources)} sources from Self-RAG winning candidate",
                )
        else:
            # Fallback: Self-RAG metadata unavailable (e.g., error path), do a fresh retrieval.
            source_docs = retrieve_quality_chunks(
                vectorstore,
                query,
                k=int(settings["rag_final_context_k"]),
                min_results=int(settings["rag_retrieval_min_results"]),
                score_threshold=float(settings["rag_retrieval_score_threshold"]),
                print_debug=print_debug,
                weight_semantic=float(
                    settings.get("weight_semantic", cfg.WEIGHT_SEMANTIC)
                ),
                weight_bm25=float(settings.get("weight_bm25", cfg.WEIGHT_BM25)),
            )

            # Format sources for display
            for i, doc in enumerate(source_docs, 1):
                doc_metadata: Dict[str, Any] = doc.metadata  # type: ignore[assignment]
                source_entry: Dict[str, Any] = {
                    "document": str(doc_metadata.get("document", "Unknown")),
                    "page": str(doc_metadata.get("page", "N/A")),
                    "content": (
                        doc.page_content[:200] + "..."
                        if len(doc.page_content) > 200
                        else doc.page_content
                    ),
                }
                sources.append(source_entry)
                if print_debug:
                    debug_log(
                        "INFO",
                        message=f"• Source {i}: {source_entry['document']} (Page {source_entry['page']})",
                    )

        if print_debug:
            print_breaker()
            debug_log("INFO", "📎", f"Processed {len(sources)} display sources for UI")

    confidence_score = None
    if getattr(st, "session_state", None) and hasattr(
        st.session_state, "self_rag_metadata"
    ):
        confidence_metrics = st.session_state.self_rag_metadata.get(
            "confidence_metrics", {}
        )
        if "total_score" in confidence_metrics:
            try:
                confidence_score = float(confidence_metrics["total_score"])
            except (ValueError, TypeError):
                pass
        elif "issup" in confidence_metrics:
            try:
                confidence_score = float(confidence_metrics["issup"])
            except (ValueError, TypeError):
                pass

    return answer_clean, sources, found_answer, reasoning_trace, confidence_score


def save_query_and_answer_to_history(
    notebook_id: str,
    query: str,
    answer: str,
    sources: List[Dict[str, Any]],
    print_debug: bool = False,
) -> None:
    """Save user query and assistant answer to chat history database."""
    from middlewares import db_middleware as db

    if print_debug:
        debug_log(
            "INFO",
            "💾",
            f"Saving query and answer to history for notebook {notebook_id}",
        )

    # Save user message
    db.add_chat_message(
        notebook_id=notebook_id,
        role=cfg.USER_ROLE_NAME,
        content=query,
    )

    # Save assistant message with sources
    db.add_chat_message(
        notebook_id=notebook_id,
        role=cfg.ASSISTANT_ROLE_NAME,
        content=answer,
        self_rag_sources=sources,
    )


def save_answer_as_note(
    notebook_id: str,
    query: str,
    answer: str,
    print_debug: bool = False,
) -> None:
    """Save an answer as a study note in the notebook."""
    from middlewares import db_middleware as db

    if print_debug:
        debug_log("INFO", "📝", f"Saving answer as note in notebook {notebook_id}")

    db.add_note(
        notebook_id=notebook_id,
        title=query,
        content=answer,
    )


# ============================================================================
# VECTORSTORE PERSISTENCE UTILITIES
# ============================================================================


def get_notebook_vectorstore_dir(notebook_id: str) -> Path:
    """Get the path to the vectorstores directory for a specific notebook."""
    base_dir = Path("data") / "vectorstores"
    base_dir.mkdir(exist_ok=True)
    nb_dir = base_dir / f"nb_{notebook_id}"
    nb_dir.mkdir(exist_ok=True)
    return nb_dir


def get_source_vectorstore_dir(notebook_id: str, source_id: str) -> Path:
    """Get the path to the isolated vectorstore for a specific source."""
    nb_dir = get_notebook_vectorstore_dir(notebook_id)
    src_dir = nb_dir / f"src_{source_id}"
    src_dir.mkdir(exist_ok=True)
    return src_dir


@st.cache_resource(show_spinner=False)
def try_load_embeddings() -> Optional[HuggingFaceEmbeddings]:
    """Try to load embeddings with GPU, fallback to CPU on OOM."""
    try:
        debug_log("INFO", "🔳", "Attempting to load embeddings on GPU...")
        return HuggingFaceEmbeddings(
            model_name=cfg.EMBEDDING_MODEL_NAME,
            model_kwargs={"device": "cuda"},
        )
    except Exception as e:
        debug_log("WARNING", message=f"GPU loading failed: {e}. Falling back to CPU...")
        try:
            return HuggingFaceEmbeddings(
                model_name=cfg.EMBEDDING_MODEL_NAME,
                model_kwargs={"device": "cpu"},
            )
        except Exception as e2:
            debug_log(
                "ERROR",
                message=f"CPU loading failed: {e2}. Embeddings cannot be loaded.",
            )
            return None


@st.cache_resource(show_spinner=False)
def try_load_cross_encoder() -> Any:
    """Try to load the Cross-Encoder model with GPU fallback to CPU."""
    try:
        from sentence_transformers import CrossEncoder
    except ImportError:
        debug_log("ERROR", message="sentence-transformers library not found.")
        return None

    try:
        debug_log("INFO", "🔳", "Attempting to load Cross-Encoder on GPU...")
        return CrossEncoder(cfg.CROSS_ENCODER_MODEL_NAME, device="cuda")
    except Exception as e:
        debug_log("WARNING", message=f"GPU loading failed: {e}. Falling back to CPU...")
        try:
            return CrossEncoder(cfg.CROSS_ENCODER_MODEL_NAME, device="cpu")
        except Exception as e2:
            debug_log(
                "ERROR",
                message=f"CPU loading failed: {e2}. Cross-Encoder cannot be loaded.",
            )
            return None


def rerank_with_cross_encoder(
    query: str,
    candidates: List[Document],
    top_k: int,
    print_debug: bool = False,
) -> List[Document]:
    """Execute two-stage retrieval reranking."""
    if len(candidates) <= top_k:
        return candidates

    import streamlit as st

    if "cross_encoder_model" not in st.session_state:
        st.session_state.cross_encoder_model = try_load_cross_encoder()

    model = st.session_state.cross_encoder_model

    if not model:
        if print_debug:
            debug_log("WARNING", "⚠️", "Cross-Encoder unavailable. Skipping reranking.")
        return candidates[:top_k]

    pairs = [[query, doc.page_content] for doc in candidates]
    if print_debug:
        debug_log(
            "INFO", "⚖️", f"Re-ranking {len(candidates)} chunks with Cross-Encoder..."
        )

    try:
        scores = model.predict(pairs, batch_size=32, show_progress_bar=False)
    except Exception as rerank_err:
        # CUDA OOM (or other GPU failure) during predict — evict the cached model so it
        # frees VRAM and will be re-initialized on the next call, then fall back to the
        # existing similarity scores so the current iteration can still proceed.
        if print_debug:
            debug_log(
                "WARNING",
                "⚠️",
                f"Cross-Encoder reranking failed ({str(rerank_err)[:100]}). "
                "Evicting cached model and falling back to similarity score order.",
            )
        # Remove the broken/OOM model from both caches so it is reloaded fresh on next call
        try_load_cross_encoder.clear()  # evict from @st.cache_resource
        if "cross_encoder_model" in st.session_state:
            del st.session_state.cross_encoder_model
        # Sort by existing FAISS L2 similarity scores (lower = closer = better)
        candidates.sort(
            key=lambda x: float(
                cast(Dict[str, Any], getattr(x, "metadata")).get(
                    "similarity_score", 999.0
                )
            )
        )
        return candidates[:top_k]

    for doc, score in zip(candidates, scores):
        cast(Dict[str, Any], getattr(doc, "metadata"))["rerank_score"] = float(score)

    candidates.sort(
        key=lambda x: float(
            cast(Dict[str, Any], getattr(x, "metadata")).get("rerank_score", 0.0)
        ),
        reverse=True,
    )
    return candidates[:top_k]


def _get_or_create_bm25_retriever(vectorstore: Any, print_debug: bool = False) -> Any:
    """Gets cached BM25 retriever or creates a new one based on vectorstore doc chunks."""
    import streamlit as st

    try:
        from langchain_community.retrievers import BM25Retriever
    except ImportError:
        if print_debug:
            debug_log("ERROR", message="BM25 package not found.")
        return None

    # Use chunk keys to represent the current corpus signature
    chunk_keys = tuple(sorted(vectorstore.docstore._dict.keys()))
    current_signature = hash(chunk_keys)

    if (
        "bm25_retriever" in st.session_state
        and "bm25_signature" in st.session_state
        and st.session_state.bm25_signature == current_signature
    ):
        return st.session_state.bm25_retriever

    if print_debug:
        debug_log("INFO", "⚙️", "Initializing BM25 Retriever from vectorstore chunks...")

    docs = list(vectorstore.docstore._dict.values())
    if not docs:
        return None

    bm25 = BM25Retriever.from_documents(docs)

    st.session_state.bm25_retriever = bm25
    st.session_state.bm25_signature = current_signature

    return bm25


def retrieve_quality_chunks(
    vectorstore: FAISS,
    query: str,
    k: int = cfg.RAG_FINAL_CONTEXT_K,
    min_results: int = cfg.RAG_RETRIEVAL_MIN_RESULTS,
    score_threshold: float = cfg.RAG_RETRIEVAL_SCORE_THRESHOLD,
    print_debug: bool = False,
    weight_semantic: float = cfg.WEIGHT_SEMANTIC,
    weight_bm25: float = cfg.WEIGHT_BM25,
) -> List[Document]:
    """
    Retrieve document chunks with quality-based filtering and fallback guarantees.

    Uses L2 distance similarity scoring to filter chunks. If strict threshold filtering
    returns fewer than min_results, falls back to returning top-K results without
    threshold filtering to ensure the LLM always gets minimum context.

    Args:
        vectorstore: FAISS vectorstore containing indexed documents.
        query: User's question to retrieve relevant chunks for.
        k: Maximum number of chunks to retrieve.
        min_results: Minimum guaranteed result count (triggers fallback if not met).
        score_threshold: L2 distance threshold; lower scores = better matches.
                        Typical range: 0.0-30.0 for multilingual embeddings.
        print_debug: Enable detailed logging of retrieval process.

    Returns:
        List[Document]: Ordered list of relevant document chunks with metadata including
                       similarity_score and passed_quality_threshold flags.

    Note:
        L2 distance metric used by FAISS: lower distance = higher similarity.
        Values typically range 4.0-6.0 for moderate semantic relevance.
    """
    SEARCH_MODE_BM25 = "BM25 Keyword Search"

    # Evaluate Gatekeeper Hybrid Configuration
    if weight_bm25 <= 0.0 or weight_semantic >= 1.0:
        search_mode = "Semantic Search"
    elif weight_semantic <= 0.0 or weight_bm25 >= 1.0:
        search_mode = SEARCH_MODE_BM25
    else:
        search_mode = "Hybrid Search"

    if print_debug:
        debug_log(
            "INFO",
            "🔎",
            f"{search_mode} | Threshold: {score_threshold} | Fallback: {min_results}",
        )

    # Step 1: Get top K results with scores
    if search_mode == "Semantic Search":
        results_with_scores: List[Any] = vectorstore.similarity_search_with_score(  # pyright: ignore[reportUnknownMemberType]
            query, k=k
        )
    else:
        bm25_retriever = _get_or_create_bm25_retriever(vectorstore, print_debug)
        if not bm25_retriever:
            debug_log(
                "WARNING",
                message="BM25 Initialization failed. Falling back to Semantic Search.",
            )
            results_with_scores = vectorstore.similarity_search_with_score(query, k=k)  # pyright: ignore[reportUnknownMemberType]
        elif search_mode == SEARCH_MODE_BM25:
            bm25_retriever.k = k
            bm25_docs = bm25_retriever.invoke(query)
            if bm25_docs:
                # Cross-check BM25 results with semantic scores so that
                # rag_retrieval_score_threshold can meaningfully filter BM25 results.
                # Without this, all BM25 docs get score=0.0 and always pass the threshold.
                sem_score_map = {
                    hash(doc.page_content[:100]): score
                    for doc, score in vectorstore.similarity_search_with_score(  # pyright: ignore[reportUnknownMemberType]
                        query, k=k * 2
                    )
                }
                results_with_scores = [
                    (
                        doc,
                        sem_score_map.get(
                            hash(doc.page_content[:100]), score_threshold + 1.0
                        ),
                    )
                    for doc in bm25_docs
                ]
            else:
                results_with_scores = []
        else:
            # Hybrid Search: fuse Semantic + BM25 via Reciprocal Rank Fusion (RRF)
            from langchain_classic.retrievers.ensemble import EnsembleRetriever
            from langchain_core.retrievers import BaseRetriever
            from langchain_core.callbacks import CallbackManagerForRetrieverRun

            class CustomFAISSRetriever(BaseRetriever):
                vectorstore: Any
                k: int
                score_threshold: float

                def _get_relevant_documents(
                    self, query: str, *, run_manager: CallbackManagerForRetrieverRun
                ):
                    results = self.vectorstore.similarity_search_with_score(
                        query, k=self.k
                    )
                    docs: List[Document] = []
                    for doc, score in results:
                        if score <= self.score_threshold:
                            docs.append(doc)
                    return docs

            sem_k = max(1, round(k * weight_semantic))
            bm25_k = max(1, round(k * weight_bm25))
            bm25_retriever.k = bm25_k
            base_retriever = CustomFAISSRetriever(
                vectorstore=vectorstore, k=sem_k, score_threshold=score_threshold
            )
            ensemble = EnsembleRetriever(
                retrievers=[base_retriever, bm25_retriever],
                weights=[weight_semantic, weight_bm25],
                c=cfg.RRF_C,
            )
            hybrid_docs = ensemble.invoke(query)

            # All hybrid docs have already been quality-filtered through CustomFAISSRetriever;
            # assign synthetic passing score so the threshold step below remains a no-op.
            results_with_scores = [(doc, 0.0) for doc in hybrid_docs]

    # Step 2: Filter by quality threshold (lower distance = better match)
    filtered_results: List[Document] = []
    passed_threshold = 0
    failed_threshold = 0

    for doc, score in results_with_scores:
        # Store similarity score in metadata for reference
        doc.metadata["similarity_score"] = score

        # Apply quality threshold: only include chunks with good similarity
        if score <= score_threshold:
            filtered_results.append(doc)
            doc.metadata["passed_quality_threshold"] = True
            passed_threshold += 1
        else:
            failed_threshold += 1

    if print_debug:
        debug_log(
            "INFO",
            "ℹ️",
            f"Quality filter: {passed_threshold} passed, {failed_threshold} failed (Total: {len(results_with_scores)})",
        )

    # Step 3: Fallback mechanism - if threshold filtered too much, use top min_results
    if len(filtered_results) < min_results:
        if print_debug:
            debug_log(
                "WARNING",
                "⚠️",
                f"Only {len(filtered_results)} chunks passed threshold (needed {min_results}). Using top {min_results} results as fallback.",
            )
        # Take the top min_results from the original sorted results
        filtered_results = []
        for doc, score in results_with_scores[:min_results]:
            filtered_results.append(doc)
            cast(Dict[str, Any], getattr(doc, "metadata"))[
                "passed_quality_threshold"
            ] = score <= score_threshold

    if search_mode == SEARCH_MODE_BM25 and len(filtered_results) == 0:
        if print_debug:
            debug_log(
                "WARNING",
                "⚠️",
                "BM25 Keyword Search returned 0 chunks. Falling back to Semantic Search (vector).",
            )

        # Fallback to Semantic Search
        sem_results = vectorstore.similarity_search_with_score(query, k=k)  # pyright: ignore[reportUnknownMemberType]

        filtered_results = []
        passed_threshold = 0
        failed_threshold = 0
        for doc, score in sem_results:
            cast(Dict[str, Any], getattr(doc, "metadata"))["similarity_score"] = score
            if score <= score_threshold:
                filtered_results.append(doc)
                cast(Dict[str, Any], getattr(doc, "metadata"))[
                    "passed_quality_threshold"
                ] = True
                passed_threshold += 1
            else:
                cast(Dict[str, Any], getattr(doc, "metadata"))[
                    "passed_quality_threshold"
                ] = False
                failed_threshold += 1

        if len(filtered_results) < min_results:
            filtered_results = []
            for doc, score in sem_results[:min_results]:
                filtered_results.append(doc)

    if search_mode == "Hybrid Search" and len(filtered_results) == 0:
        if print_debug:
            debug_log(
                "WARNING",
                "⚠️",
                "Hybrid Search returned 0 chunks. Falling back to pure Semantic Search (vector).",
            )

        # Fallback to Semantic Search
        sem_results = vectorstore.similarity_search_with_score(query, k=k)  # pyright: ignore[reportUnknownMemberType]

        filtered_results = []
        for doc, score in sem_results:
            cast(Dict[str, Any], getattr(doc, "metadata"))["similarity_score"] = score
            if score <= score_threshold:
                filtered_results.append(doc)
                cast(Dict[str, Any], getattr(doc, "metadata"))[
                    "passed_quality_threshold"
                ] = True
            else:
                cast(Dict[str, Any], getattr(doc, "metadata"))[
                    "passed_quality_threshold"
                ] = False

        if len(filtered_results) < min_results:
            filtered_results = []
            for doc, score in sem_results[:min_results]:
                filtered_results.append(doc)

    if len(filtered_results) == 0 and print_debug:
        debug_log(
            "WARNING",
            message="0 chunks passed the quality threshold and min_results. Returning empty list.",
        )

    if print_debug:
        debug_log(
            "INFO", "✅", f"Final retrieved chunks for display: {len(filtered_results)}"
        )

    return filtered_results


def format_context_with_sources(
    docs: List[Document],
    max_chunk_length: int = cfg.RAG_MAX_CHUNK_LEN,
    print_debug: bool = False,
) -> str:
    """
    Format retrieved documents into a context string with source citations.
    Intelligently limits context to prevent GPU OOM during LLM inference.

    Args:
      docs: List of Document objects (with optional similarity_score metadata)

    Returns:
      Formatted context string with page numbers and citations
    """
    if not docs:
        if print_debug:
            debug_log("INFO", "📄", "No relevant context found to format.")
        return "No relevant context found."

    context_parts: List[str] = []
    current_length = 0

    if print_debug:
        debug_log(
            "INFO",
            "📄",
            f"Formatting {len(docs)} chunks into context with max total length of {cfg.RAG_MAX_CTX_LEN} bytes...",
        )

    for idx, doc in enumerate(docs, 1):
        metadata: dict[str, Any] = getattr(doc, "metadata", {})
        page_num = str(metadata.get("page", "Unknown"))
        # Clean up the text: remove extra whitespace
        text = " ".join(doc.page_content.split())

        # Limit individual chunk to configured max length to keep relevant info focused
        if len(text) > max_chunk_length:
            text = text[:max_chunk_length] + "..."

        part = f"[Source: Page {page_num}]\n{text}"
        part_length = len(part)

        # Stop adding chunks if we exceed max total context length
        if current_length + part_length > cfg.RAG_MAX_CTX_LEN:
            if print_debug:
                debug_log(
                    "WARNING",
                    message=f"Max context length reached. Stopping at chunk {idx - 1}. (Current: {current_length}B, Limit: {cfg.RAG_MAX_CTX_LEN}B)",
                )
            break

        context_parts.append(part)
        current_length += part_length

        if print_debug:
            debug_log(
                "INFO",
                "📑",
                f"Added chunk {idx}: Page {page_num} ({part_length}B, running total: {current_length}B)",
            )

    final_context = "\n\n".join(context_parts)
    if print_debug:
        debug_log(
            "INFO",
            "📄",
            f"Final formatted context prepared with {len(context_parts)} chunks, total length: {len(final_context)} bytes",
        )

    return final_context


def reload_vectorstore_and_chain(
    notebook_id: str,
    selected_sources: Set[str],
    print_debug: bool = False,
) -> None:
    """Reload vectorstore and RAG chain based on currently selected sources."""
    st.session_state.vectorstore = load_persisted_vectorstore_filtered(
        notebook_id, selected_sources
    )

    if st.session_state.vectorstore is not None:
        from core.self_rag import (
            create_history_aware_rag_chain,
        )  # lazy import avoids circular dependency

        st.session_state.rag_chain = create_history_aware_rag_chain(
            st.session_state.vectorstore,
            print_debug=print_debug,
            notebook_id=notebook_id,
        )
    else:
        st.session_state.rag_chain = None


def load_persisted_vectorstore_filtered(
    notebook_id: str, selected_source_ids: Set[str]
) -> Optional[FAISS]:
    """Load and merge only the selected sources' vectorstores.

    Args:
        notebook_id: The notebook ID
        selected_source_ids: Set of source IDs to load

    Returns:
        Merged FAISS vectorstore or None if no valid sources selected
    """
    from middlewares import db_middleware as db

    if not selected_source_ids:
        return None

    sources = db.get_sources_for_notebook(notebook_id)
    merged_vs = None

    embeddings = try_load_embeddings()
    if not embeddings:
        return None

    for source in sources:
        source_id = source["id"]
        # Only load if source is in selected set
        if source_id not in selected_source_ids:
            continue

        vs_dir = get_source_vectorstore_dir(notebook_id, source_id)
        if (vs_dir / "index.faiss").exists():
            try:
                vs = FAISS.load_local(
                    str(vs_dir),
                    embeddings,
                    allow_dangerous_deserialization=True,
                )
                if merged_vs is None:
                    merged_vs = vs
                else:
                    merged_vs.merge_from(vs)
            except Exception as e:
                debug_log(
                    "WARNING",
                    message=f"Failed to load vectorstore for source {source_id}: {e}",
                )
    return merged_vs


def format_relative_time(dt_str: str) -> str:
    """Format a DB datetime string as a human-readable relative time."""
    try:
        dt = datetime.strptime(dt_str[:19], "%Y-%m-%d %H:%M:%S")
        now_utc = datetime.now(timezone.utc)

        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)

        secs = int((now_utc - dt).total_seconds())
        if secs < 60:
            return f"{secs}s ago"
        if secs < 3600:
            return f"{secs // 60}m ago"
        if secs < 86400:
            return f"{secs // 3600}h ago"
        return f"{secs // 86400}d ago"
    except Exception:
        return dt_str[:10]


# ============================================================================
# ADVANCED RAG: CHAT HISTORY & GREETING DETECTION
# ============================================================================


def is_greeting(user_query: str) -> bool:
    """
    Detect if user query is a greeting or general question across multiple languages.

    Supports: English, Vietnamese, and Mandarin Chinese with automatic language detection.
    Falls back to English patterns if language detection fails.
    Handles edge cases like repeated characters ("heeey" → "hey").

    Args:
        user_query: The user's input text

    Returns:
        True if query is a greeting/general question, False if it needs document search
    """
    query_lower = clean_spaces(user_query.lower())

    try:
        # Detect language (use original query for language detection)
        detected_lang: str = str(detect(query_lower))  # type: ignore[misc]

        # Map detected language codes to our supported languages
        # langdetect uses ISO 639-1 codes: 'en', 'vi', 'zh-cn', 'zh-tw', 'ko', etc.
        if detected_lang.startswith("zh"):
            lang_code = "zh"  # Both Simplified and Traditional Chinese
        elif detected_lang == "ko":
            # Korean detection can be ambiguous with Chinese for short greetings
            # Use Chinese patterns as fallback for these ambiguous cases
            lang_code = "zh"
        elif detected_lang in ("en", "vi"):
            lang_code = detected_lang
        else:
            # Unsupported language, fall back to English patterns
            lang_code = "en"

        # Get language-specific patterns
        patterns = cfg.GREETING_PATTERNS_BY_LANGUAGE.get(
            lang_code, cfg.DEFAULT_EN_GREETING_PATTERNS
        )

        # Try to match against original query first
        for pattern in patterns:
            if re.match(pattern, query_lower):
                return True

        # If no match with original query, try with normalized query (repeated characters removed)
        query_normalized = re.sub(r"(.)\1+", r"\1", query_lower)
        for pattern in patterns:
            if re.match(pattern, query_normalized):
                return True

        return False

    except Exception as e:
        # If language detection fails, fall back to English patterns
        debug_log(
            "WARNING",
            message=f"Language detection failed: {e}. Using English patterns as fallback.",
        )
        # Try original query first
        for pattern in cfg.DEFAULT_EN_GREETING_PATTERNS:
            if re.match(pattern, query_lower):
                return True

        # Try normalized query
        query_normalized = re.sub(r"(.)\1+", r"\1", query_lower)
        for pattern in cfg.DEFAULT_EN_GREETING_PATTERNS:
            if re.match(pattern, query_normalized):
                return True
        return False


def format_chat_history_for_rephrase(
    chat_history: List[Dict[str, Any]], max_messages: int = cfg.MAX_MSG_HISTORY
) -> List[BaseMessage]:
    """
    Convert database chat history to LangChain message format for history-aware retriever.

    Takes the latest N messages from chat history and converts them to LangChain's
    HumanMessage/AIMessage format for use in history-aware retrieval chains.

    Args:
        chat_history: List of message dicts from database (role, content, etc.)
        max_messages: Maximum number of messages to include (use latest N messages)

    Returns:
        List of LangChain BaseMessage objects (HumanMessage or AIMessage)
    """
    messages: List[BaseMessage] = []

    # Use only the latest max_messages to avoid context overflow
    for msg in chat_history[-max_messages:]:
        if msg["role"] == cfg.USER_ROLE_NAME:
            from langchain_core.messages import HumanMessage

            messages.append(HumanMessage(content=msg["content"]))
        elif msg["role"] == cfg.ASSISTANT_ROLE_NAME:
            from langchain_core.messages import AIMessage

            messages.append(AIMessage(content=msg["content"]))

    return messages


def format_co_rag_chat_history(
    chat_history: List[Dict[str, Any]], max_messages: int = cfg.MAX_MSG_HISTORY
) -> List[Dict[str, Any]]:
    """
    Build a Co-RAG-isolated view of the chat history.

    For assistant messages, uses ``co_rag_content`` so Co-RAG's LLM sees only
    Co-RAG's prior answers — not Self-RAG's — keeping the two pipelines fully
    independent. Falls back to ``content`` for backward-compatibility with
    messages that predate the dual-pipeline schema.

    Args:
        chat_history: Raw chat history list from DB / session state.
        max_messages: Trim to the latest N messages to avoid context overflow.

    Returns:
        List of message dicts with ``role`` and ``content`` keys suitable for
        passing into LangChain-style history formatters.
    """
    result: List[Dict[str, Any]] = []
    for msg in chat_history[-max_messages:]:
        role = msg.get("role", "")
        if role == cfg.USER_ROLE_NAME:
            result.append({"role": role, "content": msg.get("content", "")})
        elif role == cfg.ASSISTANT_ROLE_NAME:
            # Prefer Co-RAG's own answer; fall back to generic content for old messages
            co_content = msg.get("co_rag_content") or msg.get("content", "")
            result.append({"role": role, "content": co_content})
    return result


def get_installed_ollama_models() -> List[str]:
    """Fetch tags from local Ollama API (http://localhost:11434/api/tags)."""
    import requests

    try:
        response = requests.get(f"{cfg.OLLAMA_BASE_URL}/api/tags", timeout=2)
        if response.status_code == 200:
            return [m["name"] for m in response.json().get("models", [])]
        return []
    except Exception:
        return []
